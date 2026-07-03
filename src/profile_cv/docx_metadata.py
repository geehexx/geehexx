from __future__ import annotations

import os
import re
import shutil
import signal
import subprocess
import tempfile
import time
import zipfile
from collections.abc import Mapping
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Final, cast
from xml.etree import ElementTree as ET

from docx import Document

CORE_NS: Final = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS: Final = "http://purl.org/dc/elements/1.1/"
DCTERMS_NS: Final = "http://purl.org/dc/terms/"
XSI_NS: Final = "http://www.w3.org/2001/XMLSchema-instance"
APP_NS: Final = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"

ET.register_namespace("cp", CORE_NS)
ET.register_namespace("dc", DC_NS)
ET.register_namespace("dcterms", DCTERMS_NS)
ET.register_namespace("xsi", XSI_NS)
ET.register_namespace("", APP_NS)

CORE_PROPS = "docProps/core.xml"
APP_PROPS = "docProps/app.xml"
CUSTOM_PROPS = "docProps/custom.xml"
INT_APP_PROPERTIES: Final = (
    "TotalTime",
    "Pages",
    "Words",
    "Characters",
    "DocSecurity",
    "Lines",
    "Paragraphs",
    "CharactersWithSpaces",
)
BOOL_APP_PROPERTIES: Final = ("ScaleCrop", "LinksUpToDate", "SharedDoc", "HyperlinksChanged")


@dataclass(frozen=True)
class DocxStats:
    pages: int
    words: int
    characters: int
    characters_with_spaces: int
    paragraphs: int
    lines: int

    def as_app_properties(self) -> dict[str, int]:
        return {
            "Pages": self.pages,
            "Words": self.words,
            "Characters": self.characters,
            "Lines": self.lines,
            "Paragraphs": self.paragraphs,
            "CharactersWithSpaces": self.characters_with_spaces,
        }


def normalize_docx_metadata(
    docx_path: Path,
    *,
    title: str,
    subject: str,
    author: str,
    keywords: str,
    category: str = "Resume",
    description: str = "Resume for applied AI, platform, backend, and technical leadership roles.",
    modified: datetime | None = None,
    application: str = "profile-cv",
    app_version: str = "1.0.0",
) -> DocxStats:
    """Normalize shareable DOCX core and extended properties in-place."""
    docx_path = docx_path.resolve()
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)
    stats = compute_docx_stats(docx_path, render_pages=True)
    normalized_modified = _to_utc_naive(modified or datetime.now(UTC))
    with zipfile.ZipFile(docx_path) as zin:
        entries = {info.filename: zin.read(info.filename) for info in zin.infolist()}
    if CORE_PROPS not in entries or APP_PROPS not in entries:
        raise RuntimeError(f"{docx_path} is missing required DOCX metadata parts")
    entries[CORE_PROPS] = _patch_core_properties(
        entries[CORE_PROPS],
        title=title,
        subject=subject,
        author=author,
        keywords=keywords,
        category=category,
        description=description,
        modified=normalized_modified,
    )
    entries[APP_PROPS] = _patch_app_properties(
        entries[APP_PROPS], stats=stats, application=application, app_version=app_version
    )
    if CUSTOM_PROPS in entries:
        entries[CUSTOM_PROPS] = _empty_custom_properties()
    _rewrite_zip(docx_path, entries)
    return stats


def assert_docx_metadata(
    docx_path: Path,
    *,
    expected_title: str = "Andrew Crozier - Resume",
    expected_author: str = "Andrew Crozier",
    expected_application: str = "profile-cv",
) -> dict[str, int | str | bool]:
    # Recompute text statistics from the DOCX package. Page count is checked as a
    # positive value here to avoid doing a second LibreOffice conversion inside
    # every build.
    stats = compute_docx_stats(docx_path, render_pages=False)
    core = read_core_properties(docx_path)
    app = read_app_properties(docx_path)
    custom_count = count_custom_properties(docx_path)

    mismatches: list[str] = []
    app_pages = int(app.get("Pages", 0))
    if app_pages < 1:
        mismatches.append(f"Pages: expected positive rendered page count, got {app_pages}")
    for key, expected in stats.as_app_properties().items():
        if key == "Pages":
            continue
        actual = int(app.get(key, -1))
        if actual != expected:
            mismatches.append(f"{key}: expected {expected}, got {actual}")
    if core.get("title") != expected_title:
        mismatches.append(f"title: expected {expected_title!r}, got {core.get('title')!r}")
    if core.get("creator") != expected_author:
        mismatches.append(f"creator: expected {expected_author!r}, got {core.get('creator')!r}")
    if core.get("lastModifiedBy") != expected_author:
        mismatches.append(
            f"lastModifiedBy: expected {expected_author!r}, got {core.get('lastModifiedBy')!r}"
        )
    if app.get("Application") != expected_application:
        mismatches.append(
            f"Application: expected {expected_application!r}, got {app.get('Application')!r}"
        )
    if custom_count:
        mismatches.append(f"custom properties: expected 0, got {custom_count}")
    if mismatches:
        raise AssertionError(f"{docx_path} has stale DOCX metadata: {mismatches}")
    return {
        "pages": app_pages,
        "words": stats.words,
        "characters": stats.characters,
        "characters_with_spaces": stats.characters_with_spaces,
        "paragraphs": stats.paragraphs,
        "lines": stats.lines,
        "application": str(app.get("Application", "")),
        "app_version": str(app.get("AppVersion", "")),
        "title": core.get("title", ""),
        "creator": core.get("creator", ""),
        "custom_properties": custom_count,
        "metadata_current": True,
    }


def compute_docx_stats(docx_path: Path, *, render_pages: bool = True) -> DocxStats:
    document = Document(str(docx_path))
    paragraph_text = [paragraph.text for paragraph in document.paragraphs]
    text = "\n".join(paragraph_text)
    words = len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", text))
    characters = sum(1 for char in text if not char.isspace())
    characters_with_spaces = len(text)
    paragraphs = len(document.paragraphs)
    lines = max(paragraphs, len(text.splitlines()))
    pages = rendered_docx_page_count(docx_path) if render_pages else 0
    return DocxStats(
        pages=pages,
        words=words,
        characters=characters,
        characters_with_spaces=characters_with_spaces,
        paragraphs=paragraphs,
        lines=lines,
    )


def rendered_docx_page_count(docx_path: Path) -> int:
    """Return the page count after LibreOffice renders the DOCX to PDF."""
    docx_path = docx_path.resolve()
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)
    with tempfile.TemporaryDirectory(prefix="profile-cv-docx-pages-") as tmp_name:
        tmp = Path(tmp_name)
        out_dir = tmp / "out"
        home = tmp / "home"
        out_dir.mkdir()
        home.mkdir()
        env = os.environ.copy()
        env["HOME"] = str(home)
        user_installation = home.resolve().as_uri()
        last_stdout = ""
        last_stderr = ""
        last_returncode: int | None = None
        for attempt in range(1, 4):
            _wait_for_stable_file(docx_path)
            for candidate in out_dir.glob("*.pdf"):
                candidate.unlink()
            cmd = [
                "libreoffice",
                f"-env:UserInstallation={user_installation}",
                "--headless",
                "--norestore",
                "--nodefault",
                "--nofirststartwizard",
                "--nolockcheck",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(docx_path),
            ]
            try:
                proc = subprocess.Popen(  # noqa: S603 - fixed executable and controlled args.
                    cmd,
                    text=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    env=env,
                    start_new_session=True,
                )
            except FileNotFoundError as exc:
                raise RuntimeError("libreoffice is required for DOCX metadata page counts") from exc
            try:
                stdout, stderr = proc.communicate(timeout=90)
            except subprocess.TimeoutExpired as exc:
                os.killpg(proc.pid, signal.SIGTERM)
                try:
                    stdout, stderr = proc.communicate(timeout=5)
                except subprocess.TimeoutExpired:
                    os.killpg(proc.pid, signal.SIGKILL)
                    stdout, stderr = proc.communicate()
                raise RuntimeError(
                    f"LibreOffice timed out while rendering {docx_path}\n"
                    f"STDOUT:\n{stdout}\nSTDERR:\n{stderr}"
                ) from exc
            last_stdout = stdout
            last_stderr = stderr
            last_returncode = proc.returncode
            pdf_path = _wait_for_pdf(out_dir, docx_path.stem)
            if proc.returncode == 0 and pdf_path is not None:
                return _pdf_pages(pdf_path)
            if attempt < 3:
                time.sleep(0.75 * attempt)
        raise RuntimeError(
            f"LibreOffice did not render {docx_path} after 3 attempts "
            f"(last return code {last_returncode})\n"
            f"STDOUT:\n{last_stdout}\nSTDERR:\n{last_stderr}"
        )


def _wait_for_stable_file(path: Path, *, timeout_seconds: float = 5.0) -> None:
    deadline = time.monotonic() + timeout_seconds
    last_size = -1
    stable_checks = 0
    while time.monotonic() < deadline:
        if path.exists():
            size = path.stat().st_size
            if size > 0 and size == last_size:
                stable_checks += 1
                if stable_checks >= 2:
                    return
            else:
                stable_checks = 0
            last_size = size
        time.sleep(0.2)
    if not path.exists():
        raise FileNotFoundError(path)


def _wait_for_pdf(
    out_dir: Path, expected_stem: str, *, timeout_seconds: float = 5.0
) -> Path | None:
    expected = out_dir / f"{expected_stem}.pdf"
    deadline = time.monotonic() + timeout_seconds
    while time.monotonic() < deadline:
        if expected.exists():
            return expected
        candidates = list(out_dir.glob("*.pdf"))
        if candidates:
            return candidates[0]
        time.sleep(0.2)
    return expected if expected.exists() else None


def read_core_properties(docx_path: Path) -> dict[str, str]:
    with zipfile.ZipFile(docx_path) as zf:
        root = ET.fromstring(zf.read(CORE_PROPS))
    return {
        "title": _text(root, f"{{{DC_NS}}}title"),
        "subject": _text(root, f"{{{DC_NS}}}subject"),
        "creator": _text(root, f"{{{DC_NS}}}creator"),
        "keywords": _text(root, f"{{{CORE_NS}}}keywords"),
        "description": _text(root, f"{{{DC_NS}}}description"),
        "lastModifiedBy": _text(root, f"{{{CORE_NS}}}lastModifiedBy"),
        "revision": _text(root, f"{{{CORE_NS}}}revision"),
        "category": _text(root, f"{{{CORE_NS}}}category"),
        "contentStatus": _text(root, f"{{{CORE_NS}}}contentStatus"),
    }


def read_app_properties(docx_path: Path) -> dict[str, int | str | bool]:
    with zipfile.ZipFile(docx_path) as zf:
        root = ET.fromstring(zf.read(APP_PROPS))
    values: dict[str, int | str | bool] = {}
    for child in root:
        name = _local_name(child.tag)
        value = child.text or ""
        if name in INT_APP_PROPERTIES:
            values[name] = int(value)
        elif name in BOOL_APP_PROPERTIES:
            values[name] = value.lower() == "true"
        else:
            values[name] = value
    return values


def count_custom_properties(docx_path: Path) -> int:
    with zipfile.ZipFile(docx_path) as zf:
        if CUSTOM_PROPS not in zf.namelist():
            return 0
        root = ET.fromstring(zf.read(CUSTOM_PROPS))
    return sum(1 for child in root if _local_name(child.tag) == "property")


def modified_from_resume(resume: Mapping[str, object]) -> datetime:
    meta = resume.get("meta", {})
    if not isinstance(meta, Mapping):
        return datetime.now(UTC)
    raw = meta.get("lastModified")
    if not isinstance(raw, str) or not raw:
        return datetime.now(UTC)
    return _parse_datetime(raw)


def _patch_core_properties(
    xml_bytes: bytes,
    *,
    title: str,
    subject: str,
    author: str,
    keywords: str,
    category: str,
    description: str,
    modified: datetime,
) -> bytes:
    root = ET.fromstring(xml_bytes)
    _set(root, f"{{{DC_NS}}}title", title)
    _set(root, f"{{{DC_NS}}}subject", subject)
    _set(root, f"{{{DC_NS}}}creator", author)
    _set(root, f"{{{CORE_NS}}}keywords", keywords[:255])
    _set(root, f"{{{DC_NS}}}description", description)
    _set(root, f"{{{CORE_NS}}}lastModifiedBy", author)
    _set(root, f"{{{CORE_NS}}}revision", "1")
    _set(root, f"{{{CORE_NS}}}category", category)
    _set(root, f"{{{CORE_NS}}}contentStatus", "final")
    _set_datetime(root, "created", modified)
    _set_datetime(root, "modified", modified)
    return _xml_bytes(root)


def _patch_app_properties(
    xml_bytes: bytes, *, stats: DocxStats, application: str, app_version: str
) -> bytes:
    root = ET.fromstring(xml_bytes)
    _set(root, f"{{{APP_NS}}}Template", "Normal.dotm")
    _set(root, f"{{{APP_NS}}}TotalTime", "0")
    for key, value in stats.as_app_properties().items():
        _set(root, f"{{{APP_NS}}}{key}", str(value))
    _set(root, f"{{{APP_NS}}}Application", application)
    _set(root, f"{{{APP_NS}}}DocSecurity", "0")
    _set(root, f"{{{APP_NS}}}ScaleCrop", "false")
    _set(root, f"{{{APP_NS}}}LinksUpToDate", "false")
    _set(root, f"{{{APP_NS}}}SharedDoc", "false")
    _set(root, f"{{{APP_NS}}}HyperlinksChanged", "false")
    _set(root, f"{{{APP_NS}}}Company", "")
    _set(root, f"{{{APP_NS}}}AppVersion", app_version)
    return _xml_bytes(root)


def _empty_custom_properties() -> bytes:
    return (
        b'<?xml version="1.0" encoding="utf-8"?>\n'
        b"<Properties "
        b'xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" '
        b'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes" />'
    )


def _rewrite_zip(docx_path: Path, entries: Mapping[str, bytes]) -> None:
    with tempfile.NamedTemporaryFile(
        prefix=f"{docx_path.name}.", suffix=".tmp", delete=False, dir=docx_path.parent
    ) as tmp_file:
        tmp_path = Path(tmp_file.name)
    try:
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for name, content in entries.items():
                zout.writestr(name, content)
        shutil.move(str(tmp_path), docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def _pdf_pages(pdf_path: Path) -> int:
    try:
        output = subprocess.run(
            ["pdfinfo", str(pdf_path)], check=True, text=True, capture_output=True
        ).stdout
    except FileNotFoundError as exc:
        raise RuntimeError("pdfinfo is required for DOCX metadata page counts") from exc
    except subprocess.CalledProcessError as exc:
        raise RuntimeError(f"pdfinfo failed for {pdf_path}: {exc.stderr}") from exc
    for line in output.splitlines():
        if line.startswith("Pages:"):
            return int(line.split(":", 1)[1].strip())
    raise RuntimeError(f"Could not read page count from {pdf_path}")


def _set(root: ET.Element, tag: str, value: str) -> None:
    child = root.find(tag)
    if child is None:
        child = ET.SubElement(root, tag)
    child.text = value


def _set_datetime(root: ET.Element, local_name: str, value: datetime) -> None:
    tag = f"{{{DCTERMS_NS}}}{local_name}"
    child = root.find(tag)
    if child is None:
        child = ET.SubElement(root, tag)
    child.set(f"{{{XSI_NS}}}type", "dcterms:W3CDTF")
    child.text = value.strftime("%Y-%m-%dT%H:%M:%SZ")


def _text(root: ET.Element, tag: str) -> str:
    child = root.find(tag)
    if child is None or child.text is None:
        return ""
    return child.text


def _local_name(tag: str) -> str:
    return tag.split("}", 1)[-1]


def _xml_bytes(root: ET.Element) -> bytes:
    ET.indent(root, space="  ")
    return cast(bytes, ET.tostring(root, encoding="utf-8", xml_declaration=True))


def _parse_datetime(value: str) -> datetime:
    normalized = value.replace("Z", "+00:00")
    parsed = datetime.fromisoformat(normalized)
    if parsed.tzinfo is None:
        return parsed.replace(tzinfo=UTC)
    return parsed.astimezone(UTC)


def _to_utc_naive(value: datetime) -> datetime:
    if value.tzinfo is None:
        return value.replace(tzinfo=UTC)
    return value.astimezone(UTC).replace(tzinfo=None)
