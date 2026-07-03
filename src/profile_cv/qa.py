from __future__ import annotations

import shutil
import subprocess
from collections.abc import Iterable
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document

from .docx_metadata import assert_docx_metadata
from .semantic_qa import Surface, assert_semantic_alignment


@dataclass(frozen=True)
class ArtifactExpectations:
    name: str
    readme_excluded_email: str | None
    readme_sections: tuple[str, ...] = (
        "## Now",
        "## Selected Public Work",
        "## Engineering Biases",
        "## Reach Me",
    )


def expectations_from_resume(resume: dict[str, object]) -> ArtifactExpectations:
    basics = resume.get("basics", {})
    if not isinstance(basics, dict):
        return ArtifactExpectations(name="", readme_excluded_email=None)
    name = basics.get("name")
    email = basics.get("email")
    return ArtifactExpectations(
        name=str(name) if name else "",
        readme_excluded_email=str(email) if email else None,
    )


def doctor() -> dict[str, bool]:
    commands = {
        "rendercv": "rendercv",
        "pandoc": "pandoc",
        "pdfinfo": "pdfinfo",
        "pdftotext": "pdftotext",
        "libreoffice": "libreoffice",
    }
    return {name: shutil.which(command) is not None for name, command in commands.items()}


def assert_doctor(
    required: Iterable[str] = ("rendercv", "pandoc", "pdfinfo", "pdftotext", "libreoffice"),
) -> None:
    status = doctor()
    missing = [name for name in required if not status.get(name)]
    if missing:
        raise RuntimeError(f"Missing required external tools: {', '.join(missing)}")


def qa_pdf(
    pdf_path: Path,
    *,
    max_pages: int = 5,
    expectations: ArtifactExpectations | None = None,
    resume: dict[str, Any] | None = None,
) -> dict[str, int | bool]:
    if not pdf_path.exists():
        raise FileNotFoundError(pdf_path)
    text = pdftotext(pdf_path)
    pages = pdf_pages(pdf_path)
    assert_text_quality(text, source=pdf_path, expectations=expectations)
    semantic = _semantic(text, resume=resume, surface="pdf")
    if pages < 1 or pages > max_pages:
        raise AssertionError(f"Unexpected page count for {pdf_path}: {pages}")
    return {"pages": pages, "chars": len(text), "text_reviewed": True, **semantic}


def qa_docx(
    docx_path: Path,
    *,
    expectations: ArtifactExpectations | None = None,
    resume: dict[str, Any] | None = None,
) -> dict[str, int | bool]:
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)
    document = Document(str(docx_path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert_text_quality(text, source=docx_path, expectations=expectations)
    semantic = _semantic(text, resume=resume, surface="docx")
    metadata = assert_docx_metadata(docx_path)
    return {
        "paragraphs": len(document.paragraphs),
        "chars": len(text),
        "text_reviewed": True,
        "metadata_pages": int(metadata["pages"]),
        "metadata_words": int(metadata["words"]),
        "metadata_current": bool(metadata["metadata_current"]),
        **semantic,
    }


def qa_text_file(
    path: Path,
    *,
    expectations: ArtifactExpectations | None = None,
    resume: dict[str, Any] | None = None,
    surface: Surface = "markdown",
) -> dict[str, int | bool]:
    text = path.read_text(encoding="utf-8")
    assert_text_quality(text, source=path, expectations=expectations)
    semantic = _semantic(text, resume=resume, surface=surface)
    return {"chars": len(text), "text_reviewed": True, **semantic}


def qa_readme(
    path: Path,
    *,
    expectations: ArtifactExpectations | None = None,
    resume: dict[str, Any] | None = None,
) -> dict[str, int | bool]:
    text = path.read_text(encoding="utf-8")
    sections = expectations.readme_sections if expectations else ()
    for token in sections:
        if token not in text:
            raise AssertionError(f"{path} missing profile section: {token}")
    if (
        expectations
        and expectations.readme_excluded_email
        and expectations.readme_excluded_email in text
    ):
        raise AssertionError("README must not expose the direct resume email")
    expected_h1 = f"# {expectations.name}" if expectations and expectations.name else None
    if expected_h1 and text.count(expected_h1) != 1:
        raise AssertionError("README must contain exactly one H1")
    if "{{" in text or "}}" in text:
        raise AssertionError("README contains unreplaced template marker")
    semantic = _semantic(text, resume=resume, surface="readme")
    return {"chars": len(text), "text_reviewed": True, **semantic}


def assert_text_quality(
    text: str, *, source: Path, expectations: ArtifactExpectations | None = None
) -> None:
    if len(text.strip()) < 500:
        raise AssertionError(f"{source} has too little extracted text for review")
    if expectations and expectations.name and expectations.name not in text:
        raise AssertionError(f"{source} does not contain the expected profile name")


def pdftotext(pdf_path: Path) -> str:
    try:
        return subprocess.run(
            ["pdftotext", str(pdf_path), "-"], check=True, text=True, capture_output=True
        ).stdout
    except FileNotFoundError as exc:
        raise RuntimeError("pdftotext is required for PDF QA") from exc
    except subprocess.CalledProcessError as exc:
        raise RuntimeError(f"pdftotext failed for {pdf_path}: {exc.stderr}") from exc


def pdf_pages(pdf_path: Path) -> int:
    try:
        output = subprocess.run(
            ["pdfinfo", str(pdf_path)], check=True, text=True, capture_output=True
        ).stdout
    except FileNotFoundError as exc:
        raise RuntimeError("pdfinfo is required for PDF QA") from exc
    except subprocess.CalledProcessError as exc:
        raise RuntimeError(f"pdfinfo failed for {pdf_path}: {exc.stderr}") from exc
    for line in output.splitlines():
        if line.startswith("Pages:"):
            return int(line.split(":", 1)[1].strip())
    raise RuntimeError(f"Could not read page count from {pdf_path}")


def _semantic(
    text: str, *, resume: dict[str, Any] | None, surface: Surface
) -> dict[str, int | bool]:
    if resume is None:
        return {"semantic_reviewed": False}
    return assert_semantic_alignment(text, resume=resume, surface=surface).as_dict()
