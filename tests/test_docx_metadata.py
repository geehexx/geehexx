from __future__ import annotations

import zipfile
from datetime import UTC, datetime
from pathlib import Path

import pytest
from docx import Document

from profile_cv.docx_metadata import (
    assert_docx_metadata,
    normalize_docx_metadata,
    read_app_properties,
    read_core_properties,
)
from profile_cv.qa import doctor

ROOT = Path(__file__).resolve().parents[1]


def test_normalize_docx_metadata_updates_shareable_properties(tmp_path: Path) -> None:
    missing = [name for name in ("libreoffice", "java", "pdfinfo") if not doctor()[name]]
    if missing:
        pytest.skip(f"missing external tools: {missing}")

    docx_path = tmp_path / "resume.docx"
    document = Document(str(ROOT / "styles" / "reference.docx"))
    document.add_paragraph("Andrew Crozier")
    document.add_paragraph("Summary")
    document.add_paragraph("applied AI MCP platform engineering")
    document.save(str(docx_path))

    stats = normalize_docx_metadata(
        docx_path,
        title="Andrew Crozier - Resume",
        subject="Engineering Manager - Applied AI & Platform Systems",
        author="Andrew Crozier",
        keywords="resume; applied AI; MCP",
        modified=datetime(2026, 7, 1, tzinfo=UTC),
        application="profile-cv",
        app_version="1.0.0",
    )

    app = read_app_properties(docx_path)
    core = read_core_properties(docx_path)
    metadata = assert_docx_metadata(docx_path)

    assert metadata["metadata_current"] is True
    pages = app["Pages"]
    assert isinstance(pages, int)
    assert app["Application"] == "profile-cv"
    assert app["AppVersion"] == "1.0.0"
    assert app["Words"] == stats.words
    assert pages >= 1
    assert core["title"] == "Andrew Crozier - Resume"
    assert core["creator"] == "Andrew Crozier"


def test_normalize_docx_metadata_preserves_package_parts_and_removes_custom_props(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr("profile_cv.docx_metadata.rendered_docx_page_count", lambda _: 1)
    docx_path = tmp_path / "resume.docx"
    document = Document(str(ROOT / "styles" / "reference.docx"))
    document.add_paragraph("Andrew Crozier")
    document.add_paragraph("applied AI platform engineering")
    document.save(str(docx_path))
    _write_docx_parts(
        docx_path,
        {
            "customXml/item1.xml": b"<root>preserve me</root>",
            "docProps/custom.xml": (
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                b'<Properties xmlns="http://schemas.openxmlformats.org/'
                b'officeDocument/2006/custom-properties" '
                b'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/'
                b'2006/docPropsVTypes">'
                b'<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" '
                b'pid="2" name="PrivateNote"><vt:lpwstr>remove me</vt:lpwstr></property>'
                b"</Properties>"
            ),
        },
    )

    normalize_docx_metadata(
        docx_path,
        title="Andrew Crozier - Resume",
        subject="Engineering Manager - Applied AI & Platform Systems",
        author="Andrew Crozier",
        keywords="; ".join(f"keyword-{index}" for index in range(100)),
        modified=datetime(2026, 7, 1, tzinfo=UTC),
        application="profile-cv",
        app_version="1.0.0",
    )

    with zipfile.ZipFile(docx_path) as package:
        assert package.read("customXml/item1.xml") == b"<root>preserve me</root>"
    core = read_core_properties(docx_path)
    metadata = assert_docx_metadata(docx_path)
    assert len(core["keywords"]) <= 255
    assert metadata["custom_properties"] == 0


def _write_docx_parts(docx_path: Path, parts: dict[str, bytes]) -> None:
    tmp_path = docx_path.with_suffix(".tmp.docx")
    with (
        zipfile.ZipFile(docx_path) as source,
        zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as target,
    ):
        for item in source.infolist():
            if item.filename not in parts:
                target.writestr(item, source.read(item.filename))
        for filename, payload in parts.items():
            target.writestr(filename, payload)
    tmp_path.replace(docx_path)
